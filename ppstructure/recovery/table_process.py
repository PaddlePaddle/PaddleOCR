
# Copyright (c) 2022 PaddlePaddle Authors. All Rights Reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
"""
This code is refer from:https://github.com/pqzx/html2docx/blob/8f6695a778c68befb302e48ac0ed5201ddbd4524/htmldocx/h2d.py

"""
import re, argparse
import io, os
import urllib.request
from urllib.parse import urlparse
from html.parser import HTMLParser

import docx, docx.table
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from bs4 import BeautifulSoup

# values in inches
INDENT = 0.25
LIST_INDENT = 0.5
MAX_INDENT = 5.5 # To stop indents going off the page

# Style to use with tables. By default no style is used.
DEFAULT_TABLE_STYLE = None

# Style to use with paragraphs. By default no style is used.
DEFAULT_PARAGRAPH_STYLE = None


def get_filename_from_url(url):
    return os.path.basename(urlparse(url).path)

def is_url(url):
    """
    Not to be used for actually validating a url, but in our use case we only 
    care if it's a url or a file path, and they're pretty distinguishable
    """
    parts = urlparse(url)
    return all([parts.scheme, parts.netloc, parts.path])

def fetch_image(url):
    """
    Attempts to fetch an image from a url. 
    If successful returns a bytes object, else returns None
    :return:
    """
    try:
        with urllib.request.urlopen(url) as response:
            # security flaw?
            return io.BytesIO(response.read())
    except urllib.error.URLError:
        return None

def remove_last_occurence(ls, x):
    ls.pop(len(ls) - ls[::-1].index(x) - 1)

def remove_whitespace(string, leading=False, trailing=False):
    """Remove white space from a string.
    Args:
        string(str): The string to remove white space from.
        leading(bool, optional): Remove leading new lines when True.
        trailing(bool, optional): Remove trailing new lines when False.
    Returns:
        str: The input string with new line characters removed and white space squashed.
    Examples:
        Single or multiple new line characters are replaced with space.
            >>> remove_whitespace("abc\\ndef")
            'abc def'
            >>> remove_whitespace("abc\\n\\n\\ndef")
            'abc def'
        New line characters surrounded by white space are replaced with a single space.
            >>> remove_whitespace("abc \\n \\n \\n def")
            'abc def'
            >>> remove_whitespace("abc  \\n  \\n  \\n  def")
            'abc def'
        Leading and trailing new lines are replaced with a single space.
            >>> remove_whitespace("\\nabc")
            ' abc'
            >>> remove_whitespace("  \\n  abc")
            ' abc'
            >>> remove_whitespace("abc\\n")
            'abc '
            >>> remove_whitespace("abc  \\n  ")
            'abc '
        Use ``leading=True`` to remove leading new line characters, including any surrounding
        white space:
            >>> remove_whitespace("\\nabc", leading=True)
            'abc'
            >>> remove_whitespace("  \\n  abc", leading=True)
            'abc'
        Use ``trailing=True`` to remove trailing new line characters, including any surrounding
        white space:
            >>> remove_whitespace("abc  \\n  ", trailing=True)
            'abc'
    """
    # Remove any leading new line characters along with any surrounding white space
    if leading:
        string = re.sub(r'^\s*\n+\s*', '', string)

    # Remove any trailing new line characters along with any surrounding white space
    if trailing:
        string = re.sub(r'\s*\n+\s*$', '', string)

    # Replace new line characters and absorb any surrounding space.
    string = re.sub(r'\s*\n\s*', ' ', string)
    # TODO need some way to get rid of extra spaces in e.g. text <span>   </span>  text
    return re.sub(r'\s+', ' ', string)

def delete_paragraph(paragraph):
    # https://github.com/python-openxml/python-docx/issues/33#issuecomment-77661907
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

font_styles = {
    'b': 'bold',
    'strong': 'bold',
    'em': 'italic',
    'i': 'italic',
    'u': 'underline',
    's': 'strike',
    'sup': 'superscript',
    'sub': 'subscript',
    'th': 'bold',
}

font_names = {
    'code': 'Courier',
    'pre': 'Courier',
}

styles = {
    'LIST_BULLET': 'List Bullet',
    'LIST_NUMBER': 'List Number',
}

class HtmlToDocx(HTMLParser):

    def __init__(self):
        super().__init__()
        self.options = {
            'fix-html': True,
            'images': True,
            'tables': True,
            'styles': True,
        }
        self.table_row_selectors = [
            'table > tr',
            'table > thead > tr',
            'table > tbody > tr',
            'table > tfoot > tr'
        ]
        self.table_style = DEFAULT_TABLE_STYLE
        self.paragraph_style = DEFAULT_PARAGRAPH_STYLE

    def set_initial_attrs(self, document=None):
        self.tags = {
            'span': [],
            'list': [],
        }
        if document:
            self.doc = document
        else:
            self.doc = Document()
        self.bs = self.options['fix-html'] # whether or not to clean with BeautifulSoup
        self.document = self.doc
        self.include_tables = True #TODO add this option back in?
        self.include_images = self.options['images']
        self.include_styles = self.options['styles']
        self.paragraph = None
        self.skip = False
        self.skip_tag = None
        self.instances_to_skip = 0

    def copy_settings_from(self, other):
        """Copy settings from another instance of HtmlToDocx"""
        self.table_style = other.table_style
        self.paragraph_style = other.paragraph_style

    def get_cell_html(self, soup):
        # Returns string of td element with opening and closing <td> tags removed
        # Cannot use find_all as it only finds element tags and does not find text which
        # is not inside an element
        return ' '.join([str(i) for i in soup.contents])

    def add_styles_to_paragraph(self, style):
        if 'text-align' in style:
            align = style['text-align']
            if align == 'center':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == 'justify':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if 'margin-left' in style:
            margin = style['margin-left']
            units = re.sub(r'[0-9]+', '', margin)
            margin = int(float(re.sub(r'[a-z]+', '', margin)))
            if units == 'px':
                self.paragraph.paragraph_format.left_indent = Inches(min(margin // 10 * INDENT, MAX_INDENT))
            # TODO handle non px units

    def add_styles_to_run(self, style):
        if 'color' in style:
            if 'rgb' in style['color']:
                color = re.sub(r'[a-z()]+', '', style['color'])
                colors = [int(x) for x in color.split(',')]
            elif '#' in style['color']:
                color = style['color'].lstrip('#')
                colors = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
            else:
                colors = [0, 0, 0]
                # TODO map colors to named colors (and extended colors...)
                # For now set color to black to prevent crashing
            self.run.font.color.rgb = RGBColor(*colors)
            
        if 'background-color' in style:
            if 'rgb' in style['background-color']:
                color = color = re.sub(r'[a-z()]+', '', style['background-color'])
                colors = [int(x) for x in color.split(',')]
            elif '#' in style['background-color']:
                color = style['background-color'].lstrip('#')
                colors = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
            else:
                colors = [0, 0, 0]
                # TODO map colors to named colors (and extended colors...)
                # For now set color to black to prevent crashing
            self.run.font.highlight_color = WD_COLOR.GRAY_25 #TODO: map colors

    def apply_paragraph_style(self, style=None):
        try:
            if style:
                self.paragraph.style = style
            elif self.paragraph_style:
                self.paragraph.style = self.paragraph_style
        except KeyError as e:
            raise ValueError(f"Unable to apply style {self.paragraph_style}.") from e

    def parse_dict_string(self, string, separator=';'):
        new_string = string.replace(" ", '').split(separator)
        string_dict = dict([x.split(':') for x in new_string if ':' in x])
        return string_dict

    def handle_li(self):
        # check list stack to determine style and depth
        list_depth = len(self.tags['list'])
        if list_depth:
            list_type = self.tags['list'][-1]
        else:
            list_type = 'ul' # assign unordered if no tag

        if list_type == 'ol':
            list_style = styles['LIST_NUMBER']
        else:
            list_style = styles['LIST_BULLET']

        self.paragraph = self.doc.add_paragraph(style=list_style)            
        self.paragraph.paragraph_format.left_indent = Inches(min(list_depth * LIST_INDENT, MAX_INDENT))
        self.paragraph.paragraph_format.line_spacing = 1

    def add_image_to_cell(self, cell, image):
        # python-docx doesn't have method yet for adding images to table cells. For now we use this
        paragraph = cell.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image)

    def handle_img(self, current_attrs):
        if not self.include_images:
            self.skip = True
            self.skip_tag = 'img'
            return
        src = current_attrs['src']
        # fetch image
        src_is_url = is_url(src)
        if src_is_url:
            try:
                image = fetch_image(src)
            except urllib.error.URLError:
                image = None
        else:
            image = src
        # add image to doc
        if image:
            try:
                if isinstance(self.doc, docx.document.Document):
                    self.doc.add_picture(image)
                else:
                    self.add_image_to_cell(self.doc, image)
            except FileNotFoundError:
                image = None
        if not image:
            if src_is_url:
                self.doc.add_paragraph("<image: %s>" % src)
            else:
                # avoid exposing filepaths in document
                self.doc.add_paragraph("<image: %s>" % get_filename_from_url(src))
        

    def handle_table(self, html):
        """
        To handle nested tables, we will parse tables manually as follows:
        Get table soup
        Create docx table
        Iterate over soup and fill docx table with new instances of this parser
        Tell HTMLParser to ignore any tags until the corresponding closing table tag
        """
        doc = Document()
        table_soup = BeautifulSoup(html, 'html.parser')
        rows, cols_len = self.get_table_dimensions(table_soup)
        table = doc.add_table(len(rows), cols_len)
        table.style = doc.styles['Table Grid']
        cell_row = 0
        for index, row in enumerate(rows):
            cols = self.get_table_columns(row)
            cell_col = 0
            for col in cols:
                colspan = int(col.attrs.get('colspan', 1))
                rowspan = int(col.attrs.get('rowspan', 1))

                cell_html = self.get_cell_html(col)
                
                if col.name == 'th':
                    cell_html = "<b>%s</b>" % cell_html
                docx_cell = table.cell(cell_row, cell_col)
                while docx_cell.text != '':  # Skip the merged cell
                    cell_col += 1
                    docx_cell = table.cell(cell_row, cell_col)

                cell_to_merge = table.cell(cell_row + rowspan - 1, cell_col + colspan - 1)
                if docx_cell != cell_to_merge:
                    docx_cell.merge(cell_to_merge)

                child_parser = HtmlToDocx()
                child_parser.copy_settings_from(self)

                child_parser.add_html_to_cell(cell_html or ' ', docx_cell)  # occupy the position

                cell_col += colspan
            cell_row += 1
        
        # skip all tags until corresponding closing tag
        self.instances_to_skip = len(table_soup.find_all('table'))
        self.skip_tag = 'table'
        self.skip = True
        self.table = None
        return table

    def handle_link(self, href, text):
        # Link requires a relationship
        is_external = href.startswith('http')
        rel_id = self.paragraph.part.relate_to(
            href,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True  # don't support anchor links for this library yet
        )

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), rel_id)


        # Create sub-run
        subrun = self.paragraph.add_run()
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # add default color
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), "0000EE")
        rPr.append(c)

        # add underline
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)

        subrun._r.append(rPr)
        subrun._r.text = text

        # Add subrun to hyperlink
        hyperlink.append(subrun._r)

        # Add hyperlink to run
        self.paragraph._p.append(hyperlink)

    def handle_starttag(self, tag, attrs):
        if self.skip:
            return
        if tag == 'head':
            self.skip = True
            self.skip_tag = tag
            self.instances_to_skip = 0
            return
        elif tag == 'body':
            return

        current_attrs = dict(attrs)

        if tag == 'span':
            self.tags['span'].append(current_attrs)
            return
        elif tag == 'ol' or tag == 'ul':
            self.tags['list'].append(tag)
            return # don't apply styles for now
        elif tag == 'br':
            self.run.add_break()
            return

        self.tags[tag] = current_attrs
        if tag in ['p', 'pre']:
            self.paragraph = self.doc.add_paragraph()
            self.apply_paragraph_style()

        elif tag == 'li':
            self.handle_li()

        elif tag == "hr":

            # This implementation was taken from:
            # https://github.com/python-openxml/python-docx/issues/105#issuecomment-62806373

            self.paragraph = self.doc.add_paragraph()
            pPr = self.paragraph._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            pPr.insert_element_before(pBdr,
                'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                'w:pPrChange'
            )
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'auto')
            pBdr.append(bottom)

        elif re.match('h[1-9]', tag):
            if isinstance(self.doc, docx.document.Document):
                h_size = int(tag[1])
                self.paragraph = self.doc.add_heading(level=min(h_size, 9))
            else:
                self.paragraph = self.doc.add_paragraph()

        elif tag == 'img':
            self.handle_img(current_attrs)
            return

        elif tag == 'table':
            self.handle_table()
            return

        # set new run reference point in case of leading line breaks
        if tag in ['p', 'li', 'pre']:
            self.run = self.paragraph.add_run()

        # add style
        if not self.include_styles:
            return
        if 'style' in current_attrs and self.paragraph:
            style = self.parse_dict_string(current_attrs['style'])
            self.add_styles_to_paragraph(style)

    def handle_endtag(self, tag):
        if self.skip:
            if not tag == self.skip_tag:
                return

            if self.instances_to_skip > 0:
                self.instances_to_skip -= 1
                return

            self.skip = False
            self.skip_tag = None
            self.paragraph = None

        if tag == 'span':
            if self.tags['span']:
                self.tags['span'].pop()
                return
        elif tag == 'ol' or tag == 'ul':
            remove_last_occurence(self.tags['list'], tag)
            return
        elif tag == 'table':
            self.table_no += 1
            self.table = None
            self.doc = self.document
            self.paragraph = None

        if tag in self.tags:
            self.tags.pop(tag)
        # maybe set relevant reference to None?

    def handle_data(self, data):
        if self.skip:
            return

        # Only remove white space if we're not in a pre block.
        if 'pre' not in self.tags:
            # remove leading and trailing whitespace in all instances
            data = remove_whitespace(data, True, True)

        if not self.paragraph:
            self.paragraph = self.doc.add_paragraph()
            self.apply_paragraph_style()

        # There can only be one nested link in a valid html document
        # You cannot have interactive content in an A tag, this includes links
        # https://html.spec.whatwg.org/#interactive-content
        link = self.tags.get('a')
        if link:
            self.handle_link(link['href'], data)
        else:
            # If there's a link, dont put the data directly in the run
            self.run = self.paragraph.add_run(data)
            spans = self.tags['span']
            for span in spans:
                if 'style' in span:
                    style = self.parse_dict_string(span['style'])
                    self.add_styles_to_run(style)

            # add font style and name
            for tag in self.tags:
                if tag in font_styles:
                    font_style = font_styles[tag]
                    setattr(self.run.font, font_style, True)

                if tag in font_names:
                    font_name = font_names[tag]
                    self.run.font.name = font_name

    def ignore_nested_tables(self, tables_soup):
        """
        Returns array containing only the highest level tables
        Operates on the assumption that bs4 returns child elements immediately after
        the parent element in `find_all`. If this changes in the future, this method will need to be updated
        :return:
        """
        new_tables = []
        nest = 0
        for table in tables_soup:
            if nest:
                nest -= 1
                continue
            new_tables.append(table)
            nest = len(table.find_all('table'))
        return new_tables

    def get_table_rows(self, table_soup):
        # If there's a header, body, footer or direct child tr tags, add row dimensions from there
        return table_soup.select(', '.join(self.table_row_selectors), recursive=False)

    def get_table_columns(self, row):
        # Get all columns for the specified row tag.
        return row.find_all(['th', 'td'], recursive=False) if row else []

    def get_table_dimensions(self, table_soup):
        # Get rows for the table
        rows = self.get_table_rows(table_soup)
        # Table is either empty or has non-direct children between table and tr tags
        # Thus the row dimensions and column dimensions are assumed to be 0

        cols = self.get_table_columns(rows[0]) if rows else []
        # Add colspan calculation column number
        col_count = 0
        for col in cols:
            colspan = col.attrs.get('colspan', 1)
            col_count += int(colspan)

        # return len(rows), col_count
        return rows, col_count

    def get_tables(self):
        if not hasattr(self, 'soup'):
            self.include_tables = False
            return
            # find other way to do it, or require this dependency?
        self.tables = self.ignore_nested_tables(self.soup.find_all('table'))  
        self.table_no = 0

    def run_process(self, html):
        if self.bs and BeautifulSoup:
            self.soup = BeautifulSoup(html, 'html.parser')
            html = str(self.soup)
        if self.include_tables:
            self.get_tables()
        self.feed(html)

    def add_html_to_document(self, html, document):
        if not isinstance(html, str):
            raise ValueError('First argument needs to be a %s' % str)
        elif not isinstance(document, docx.document.Document) and not isinstance(document, docx.table._Cell):
            raise ValueError('Second argument needs to be a %s' % docx.document.Document)
        self.set_initial_attrs(document)
        self.run_process(html)

    def add_html_to_cell(self, html, cell):
        self.set_initial_attrs(cell)
        self.run_process(html)

    def parse_html_file(self, filename_html, filename_docx=None):
        with open(filename_html, 'r') as infile:
            html = infile.read()
        self.set_initial_attrs()
        self.run_process(html)
        if not filename_docx:
            path, filename = os.path.split(filename_html)
            filename_docx = '%s/new_docx_file_%s' % (path, filename)
        self.doc.save('%s.docx' % filename_docx)
    
    def parse_html_string(self, html):
        self.set_initial_attrs()
        self.run_process(html)
        return self.doc