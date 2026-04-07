# Copyright (c) 2025 PaddlePaddle Authors. All Rights Reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
import re
from collections import Counter
from pathlib import Path

from ..base import BaseConverter, ConvertResult
from ..registry import default_registry

# Regex patterns for Chinese numbered headings
_RE_H2 = re.compile(r"^[一二三四五六七八九十百千]+[、．.]")
_RE_H3 = re.compile(r"^（[一二三四五六七八九十百千]+）")

# Regex for field-code hyperlink instruction
_RE_FIELD_HYPERLINK = re.compile(r'HYPERLINK\s+"([^"]+)"')

# Word XML namespace
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W = "{" + _W_NS + "}"

# OMML math namespace
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_M = "{" + _M_NS + "}"


def _escape_md_url(url: str) -> str:
    """Escape parentheses in URL for Markdown link syntax."""
    return url.replace("(", "%28").replace(")", "%29")


def _convert_omath(omath_element) -> str:
    """Convert an m:oMath lxml element to LaTeX string. Returns empty string on failure."""
    try:
        from ..math.omml import oMath2Latex

        return str(oMath2Latex(omath_element)).strip()
    except Exception:
        return ""


def _paragraph_has_math(para) -> bool:
    """Check if paragraph XML contains OMML math elements."""
    return para._element.find(f".//{_M}oMath") is not None


def _paragraph_math_to_markdown(para) -> str:
    """Convert a paragraph containing OMML math to Markdown.

    Handles mixed content: text runs, hyperlinks, inline math (m:oMath),
    and display math (m:oMathPara).
    """
    text_items = []
    parts = []

    def flush_text():
        if text_items:
            md = _runs_to_markdown(text_items)
            if md:
                parts.append(md)
            text_items.clear()

    for child in para._element:
        tag = child.tag
        local = tag.split("}")[-1] if "}" in tag else tag

        if tag == f"{_M}oMathPara":
            flush_text()
            # Find inner m:oMath elements
            for omath in child.findall(f"{_M}oMath"):
                latex = _convert_omath(omath)
                if latex:
                    parts.append(f"$$\n{latex}\n$$")
        elif tag == f"{_M}oMath":
            flush_text()
            latex = _convert_omath(child)
            if latex:
                parts.append(f"${latex}$")
        elif local == "r":
            # Plain run
            try:
                from docx.text.run import Run

                run = Run(child, para)
                if run.text:
                    text_items.append(
                        (
                            bool(run.bold),
                            bool(run.italic),
                            bool(run.underline),
                            bool(run.font.strike),
                            run.text,
                            "",
                        )
                    )
            except Exception:
                pass
        elif local == "hyperlink":
            try:
                from docx.text.hyperlink import Hyperlink

                hl = Hyperlink(child, para)
                try:
                    url = hl.url or ""
                except (KeyError, AttributeError):
                    url = ""
                for run in hl.runs:
                    if run.text:
                        text_items.append(
                            (
                                bool(run.bold),
                                bool(run.italic),
                                False,
                                bool(run.font.strike),
                                run.text,
                                url,
                            )
                        )
            except Exception:
                pass

    flush_text()
    return "".join(parts)


def _paragraph_math_to_html(para) -> str:
    """Convert a paragraph containing OMML math to HTML inline text."""
    text_items = []
    parts = []

    def flush_text():
        if text_items:
            html = _runs_to_html(text_items)
            if html:
                parts.append(html)
            text_items.clear()

    for child in para._element:
        tag = child.tag
        local = tag.split("}")[-1] if "}" in tag else tag

        if tag == f"{_M}oMathPara":
            flush_text()
            for omath in child.findall(f"{_M}oMath"):
                latex = _convert_omath(omath)
                if latex:
                    parts.append(f"$$\n{latex}\n$$")
        elif tag == f"{_M}oMath":
            flush_text()
            latex = _convert_omath(child)
            if latex:
                parts.append(f"${latex}$")
        elif local == "r":
            try:
                from docx.text.run import Run

                run = Run(child, para)
                if run.text:
                    text_items.append(
                        (
                            bool(run.bold),
                            bool(run.italic),
                            bool(run.underline),
                            bool(run.font.strike),
                            run.text,
                            "",
                        )
                    )
            except Exception:
                pass
        elif local == "hyperlink":
            try:
                from docx.text.hyperlink import Hyperlink

                hl = Hyperlink(child, para)
                try:
                    url = hl.url or ""
                except (KeyError, AttributeError):
                    url = ""
                for run in hl.runs:
                    if run.text:
                        text_items.append(
                            (
                                bool(run.bold),
                                bool(run.italic),
                                False,
                                bool(run.font.strike),
                                run.text,
                                url,
                            )
                        )
            except Exception:
                pass

    flush_text()
    return "".join(parts)


def _get_body_font_size(doc) -> float:
    """Return the most common font size in the document (used as body size). Defaults to 16.0."""
    sizes: Counter = Counter()
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        for run in p.runs:
            if run.font.size:
                sizes[run.font.size.pt] += 1
                break  # only check the first run with an explicit size per paragraph
    return sizes.most_common(1)[0][0] if sizes else 16.0


def _detect_heading_level(para, body_font_size: float) -> int:
    """Return heading level (0 = not a heading, 1-6 = heading level)."""
    # Prefer Word built-in Heading styles
    style_name = para.style.name if para.style else ""
    if style_name.startswith("Heading"):
        try:
            return int(style_name.split()[-1])
        except ValueError:
            return 1

    text = para.text.strip()
    if not text:
        return 0

    # Use font size of the first run that has an explicit size
    font_size = None
    for run in para.runs:
        if run.font.size:
            font_size = run.font.size.pt
            break

    # Significantly larger than body font -> treat as heading (threshold: 1.5x, short paragraphs only)
    if font_size and font_size > body_font_size * 1.5:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return 1
        if len(text) <= 60:
            return 2

    # Chinese numbered heading patterns
    if _RE_H2.match(text):
        return 2
    if _RE_H3.match(text):
        return 3

    return 0


def _parse_field_hyperlinks(para) -> dict:
    """Parse w:fldChar field-code hyperlinks in a paragraph.

    Returns {id(run_element): url} for runs that are display text of a HYPERLINK field.
    Handles nested fields (e.g. PAGEREF inside a hyperlink result phase) via depth counter.
    """
    field_urls: dict[int, str] = {}
    phase = None  # None | "instr" | "result"
    field_url = None
    nest_depth = 0  # depth of nested fields inside result phase

    for child in para._element:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag != "r":
            continue
        fld_char = child.find(f"{_W}fldChar")
        if fld_char is not None:
            fld_type = fld_char.get(f"{_W}fldCharType")
            if fld_type == "begin":
                if phase == "result":
                    # Nested field (e.g. PAGEREF), increase depth and skip
                    nest_depth += 1
                else:
                    phase = "instr"
                    field_url = None
            elif fld_type == "separate":
                if nest_depth == 0:
                    phase = "result"
            elif fld_type == "end":
                if nest_depth > 0:
                    nest_depth -= 1
                else:
                    phase = None
                    field_url = None
            continue

        instr_elem = child.find(f"{_W}instrText")
        if instr_elem is not None and phase == "instr" and nest_depth == 0:
            if instr_elem.text:
                m = _RE_FIELD_HYPERLINK.search(instr_elem.text)
                if m:
                    field_url = m.group(1)
            continue

        if phase == "result" and nest_depth == 0 and field_url:
            t_elem = child.find(f"{_W}t")
            if t_elem is not None and t_elem.text:
                field_urls[id(child)] = field_url

    return field_urls


def _iter_paragraph_items(para) -> list:
    """Extract (bold, italic, underline, strikethrough, text, url) tuples from a paragraph in document order.

    Handles python-docx Hyperlink objects and w:fldChar field-code hyperlinks.
    Silently degrades to plain text on error.
    Note: underline is forced to False inside Hyperlink runs to avoid Word's default hyperlink underline style.
    """
    try:
        from docx.text.hyperlink import Hyperlink
    except ImportError:
        return [
            (
                bool(r.bold),
                bool(r.italic),
                bool(r.underline),
                bool(r.font.strike),
                r.text,
                "",
            )
            for r in para.runs
            if r.text
        ]

    try:
        field_urls = _parse_field_hyperlinks(para)
    except Exception:
        field_urls = {}

    items = []
    try:
        content_iter = para.iter_inner_content()
    except Exception:
        return [
            (
                bool(r.bold),
                bool(r.italic),
                bool(r.underline),
                bool(r.font.strike),
                r.text,
                "",
            )
            for r in para.runs
            if r.text
        ]

    for element in content_iter:
        try:
            if isinstance(element, Hyperlink):
                try:
                    url = element.url or ""
                except (KeyError, AttributeError):
                    url = ""
                for run in element.runs:
                    if not run.text:
                        continue
                    # Force underline=False: Word's Hyperlink style adds underline by default
                    items.append(
                        (
                            bool(run.bold),
                            bool(run.italic),
                            False,
                            bool(run.font.strike),
                            run.text,
                            url,
                        )
                    )
                # Fallback: hyperlink with no runs but has text
                if not element.runs and element.text:
                    items.append((False, False, False, False, element.text, url))
            else:
                # Plain Run
                if not element.text:
                    continue
                url = field_urls.get(id(element._element), "")
                items.append(
                    (
                        bool(element.bold),
                        bool(element.italic),
                        bool(element.underline),
                        bool(element.font.strike),
                        element.text,
                        url,
                    )
                )
        except Exception:
            continue

    return items


def _merge_runs(items) -> list:
    """Merge adjacent items with identical (bold, italic, underline, strikethrough, url).

    Returns [(bold, italic, underline, strikethrough, text, url)].
    """
    merged: list[tuple[bool, bool, bool, bool, str, str]] = []
    for bold, italic, underline, strikethrough, text, url in items:
        if not text:
            continue
        if (
            merged
            and merged[-1][0] == bold
            and merged[-1][1] == italic
            and merged[-1][2] == underline
            and merged[-1][3] == strikethrough
            and merged[-1][5] == url
        ):
            merged[-1] = (
                bold,
                italic,
                underline,
                strikethrough,
                merged[-1][4] + text,
                url,
            )
        else:
            merged.append((bold, italic, underline, strikethrough, text, url))
    return merged


def _runs_to_markdown(items) -> str:
    """Convert paragraph items to Markdown inline text, merging adjacent items with identical formatting."""
    parts = []
    for bold, italic, underline, strikethrough, text, url in _merge_runs(items):
        if bold or italic or underline or strikethrough:
            # CommonMark: marker characters must not be surrounded by spaces
            leading = len(text) - len(text.lstrip())
            trailing = len(text) - len(text.rstrip())
            prefix = text[:leading] if leading else ""
            suffix = text[len(text) - trailing :] if trailing else ""
            inner = text.strip()
            if inner:
                # Apply strikethrough first (innermost)
                if strikethrough:
                    inner = f"~~{inner}~~"
                # Apply bold/italic
                if bold and italic:
                    inner = f"***{inner}***"
                elif bold:
                    inner = f"**{inner}**"
                elif italic:
                    inner = f"*{inner}*"
                # Apply underline last (outermost)
                if underline:
                    inner = f"<u>{inner}</u>"
                text = prefix + inner + suffix
        if url:
            text = f"[{text}]({_escape_md_url(url)})"
        parts.append(text)
    # Prevent bold/italic/strikethrough markers from merging with adjacent alphanumeric text (CommonMark requirement)
    result = []
    for i, part in enumerate(parts):
        if i > 0 and result:
            prev = result[-1]
            # Previous part ends with closing marker and current part starts with alphanumeric
            if prev.endswith(("**", "*", "~~")) and part and part[0].isalnum():
                result.append("\u200b")
        result.append(part)
    return "".join(result)


def _runs_to_html(items) -> str:
    """Convert paragraph items to HTML inline text."""
    parts = []
    for bold, italic, underline, strikethrough, text, url in _merge_runs(items):
        if bold:
            text = f"<b>{text}</b>"
        if italic:
            text = f"<i>{text}</i>"
        if underline:
            text = f"<u>{text}</u>"
        if strikethrough:
            text = f"<del>{text}</del>"
        if url:
            text = f'<a href="{url}">{text}</a>'
        parts.append(text)
    return "".join(parts)


_CODE_FONTS = {
    "Courier New",
    "Courier",
    "Consolas",
    "Monaco",
    "Menlo",
    "Source Code Pro",
    "Fira Code",
    "DejaVu Sans Mono",
    "monospace",
}


def _is_code_paragraph(para) -> bool:
    """Return True if all text-bearing runs in the paragraph use a monospace font."""
    runs_with_text = [r for r in para.runs if r.text.strip()]
    if not runs_with_text:
        return False
    runs_with_font = [r for r in runs_with_text if r.font.name]
    # At least one run must have an explicit font, and all such runs must be monospace
    if not runs_with_font:
        return False
    return all(r.font.name in _CODE_FONTS for r in runs_with_font)


def _get_content_width(doc) -> int:
    """Return the content area width of the document in EMU."""
    section = doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def _table_to_html(
    table, doc, image_counter: list, images: dict, content_width: int = 0
) -> str:
    """Convert a python-docx Table to an HTML table, handling merged cells and inline images."""
    grid = [[cell for cell in row.cells] for row in table.rows]
    nrows = len(grid)
    if nrows == 0:
        return ""
    ncols = len(grid[0])

    visited: set[tuple[int, int]] = set()
    html_parts = ["<table>"]

    for i in range(nrows):
        html_parts.append("<tr>")
        for j in range(ncols):
            if (i, j) in visited:
                continue
            tc = grid[i][j]._tc
            # Compute colspan
            colspan = 1
            while j + colspan < ncols and grid[i][j + colspan]._tc is tc:
                visited.add((i, j + colspan))
                colspan += 1
            # Compute rowspan
            rowspan = 1
            while i + rowspan < nrows and grid[i + rowspan][j]._tc is tc:
                for k in range(colspan):
                    visited.add((i + rowspan, j + k))
                rowspan += 1

            cell = grid[i][j]
            content_parts = []
            for para in cell.paragraphs:
                img_list = _extract_images_from_paragraph(para, doc, image_counter)
                for filename, img_bytes, cx_emu in img_list:
                    rel_path = f"images/{filename}"
                    images[rel_path] = img_bytes
                    if cx_emu and content_width:
                        pct = min(round(cx_emu / content_width * 100), 100)
                        content_parts.append(
                            f'<img src="images/{filename}" width="{pct}%">'
                        )
                    else:
                        content_parts.append(f'<img src="images/{filename}">')
                para_html = (
                    _paragraph_math_to_html(para).strip()
                    if _paragraph_has_math(para)
                    else _runs_to_html(_iter_paragraph_items(para)).strip()
                    or para.text.strip()
                )
                if para_html:
                    content_parts.append(para_html)
            cell_html = "<br>".join(content_parts) if content_parts else ""

            tag = "th" if i == 0 else "td"
            attrs = ""
            if colspan > 1:
                attrs += f' colspan="{colspan}"'
            if rowspan > 1:
                attrs += f' rowspan="{rowspan}"'
            html_parts.append(f"<{tag}{attrs}>{cell_html}</{tag}>")
        html_parts.append("</tr>")

    html_parts.append("</table>")
    return "\n".join(html_parts)


_MIME_TO_EXT = {
    "image/jpeg": "jpeg",
    "image/jpg": "jpg",
    "image/png": "png",
    "image/gif": "gif",
    "image/bmp": "bmp",
    "image/tiff": "tiff",
    "image/webp": "webp",
    "image/svg+xml": "svg",
    "image/x-emf": "emf",
    "image/x-wmf": "wmf",
}


def _extract_images_from_paragraph(para, doc, image_counter: list) -> list:
    """Extract images from a paragraph. Returns [(filename, bytes, cx_emu)] where cx_emu is 0 if unknown."""
    WP_NS = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
    A_NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
    R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"

    results = []
    containers = para._element.findall(f".//{WP_NS}inline") + para._element.findall(
        f".//{WP_NS}anchor"
    )
    for container in containers:
        extent = container.find(f"{WP_NS}extent")
        cx_emu = int(extent.get("cx", 0)) if extent is not None else 0

        blip = container.find(f".//{A_NS}blip")
        if blip is None:
            continue
        r_embed = blip.get(f"{R_NS}embed")
        if not r_embed:
            continue
        try:
            rel = doc.part.rels[r_embed]
            img_bytes = rel.target_part.blob
            content_type = rel.target_part.content_type
            ext = _MIME_TO_EXT.get(content_type, "png")
            image_counter[0] += 1
            filename = f"image{image_counter[0]}.{ext}"
            results.append((filename, img_bytes, cx_emu))
        except (KeyError, AttributeError):
            pass
    return results


def _build_numbering_map(doc) -> dict:
    """Parse numbering.xml and return {numId: {ilvl: numFmt}} mapping."""
    WP = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    numbering_map = {}
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        return numbering_map
    numbering_elem = numbering_part._element

    # Build abstractNumId -> {ilvl: numFmt}
    abstract = {}
    for abs_num in numbering_elem.findall(f"{WP}abstractNum"):
        abs_id = abs_num.get(f"{WP}abstractNumId")
        levels = {}
        for lvl in abs_num.findall(f"{WP}lvl"):
            ilvl = int(lvl.get(f"{WP}ilvl", "0"))
            fmt_elem = lvl.find(f"{WP}numFmt")
            fmt = (
                fmt_elem.get(f"{WP}val", "bullet") if fmt_elem is not None else "bullet"
            )
            levels[ilvl] = fmt
        abstract[abs_id] = levels

    # Map numId -> abstractNumId
    for num in numbering_elem.findall(f"{WP}num"):
        num_id = num.get(f"{WP}numId")
        abs_ref = num.find(f"{WP}abstractNumId")
        if abs_ref is not None:
            abs_id = abs_ref.get(f"{WP}val")
            if abs_id in abstract:
                numbering_map[num_id] = abstract[abs_id]
    return numbering_map


def _get_list_info(para, numbering_map) -> tuple | None:
    """Return (list_type, ilvl, num_id) or None. list_type: 'bullet' | 'ordered'"""
    WP = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    pPr = para._element.find(f"{WP}pPr")
    if pPr is None:
        return None
    numPr = pPr.find(f"{WP}numPr")
    if numPr is None:
        return None
    numId_elem = numPr.find(f"{WP}numId")
    ilvl_elem = numPr.find(f"{WP}ilvl")
    if numId_elem is None:
        return None
    num_id = numId_elem.get(f"{WP}val")
    ilvl = int(ilvl_elem.get(f"{WP}val", "0")) if ilvl_elem is not None else 0
    if num_id not in numbering_map:
        return None
    fmt = numbering_map[num_id].get(ilvl, "bullet")
    list_type = (
        "ordered"
        if fmt in ("decimal", "lowerLetter", "upperLetter", "lowerRoman", "upperRoman")
        else "bullet"
    )
    return (list_type, ilvl, num_id)


def _convert_body(doc) -> tuple:
    """Traverse body elements in order and produce Markdown. Returns (markdown_str, images_dict)."""
    try:
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError:
        raise RuntimeError(
            "DOCX conversion requires python-docx: pip install paddleocr[doc2md]"
        )

    body_font_size = _get_body_font_size(doc)
    content_width = _get_content_width(doc)
    numbering_map = _build_numbering_map(doc)
    lines: list[str] = []
    images: dict = {}
    image_counter = [0]  # wrapped in list so inner functions can mutate it
    code_buf: list[str] = []  # buffer for consecutive code paragraphs
    ol_counters: dict[str, int] = {}  # key = "{numId}-{ilvl}", value = current index
    prev_was_list = False

    def flush_code_buf():
        """Flush the code buffer as a fenced code block."""
        if code_buf:
            lines.append("```")
            lines.extend(code_buf)
            lines.append("```")
            lines.append("")
            code_buf.clear()

    for child in doc.element.body:
        tag = child.tag.split("}")[-1]

        if tag == "p":
            para = Paragraph(child, doc)

            # Extract images first
            img_list = _extract_images_from_paragraph(para, doc, image_counter)
            for filename, img_bytes, cx_emu in img_list:
                flush_code_buf()
                rel_path = f"images/{filename}"
                images[rel_path] = img_bytes
                if cx_emu and content_width:
                    pct = min(round(cx_emu / content_width * 100), 100)
                    lines.append(f'<img src="images/{filename}" width="{pct}%">')
                else:
                    lines.append(f'<img src="images/{filename}">')
                lines.append("")

            text = para.text.strip()

            # Math formula detection — must check before skipping empty-text paragraphs
            # (pure formula paragraphs have para.text == "")
            if _paragraph_has_math(para):
                flush_code_buf()
                if prev_was_list:
                    lines.append("")
                prev_was_list = False
                math_md = _paragraph_math_to_markdown(para)
                if math_md:
                    lines.append(math_md)
                    lines.append("")
                continue

            if not text:
                if not img_list:
                    if code_buf:
                        code_buf.append("")  # preserve blank lines inside code blocks
                    elif lines and lines[-1] != "":
                        lines.append("")
                continue

            # Code paragraph: buffer it without heading/inline formatting
            if _is_code_paragraph(para):
                code_buf.append(para.text)
                continue

            # Non-code paragraph: flush any buffered code first
            flush_code_buf()

            level = _detect_heading_level(para, body_font_size)
            inline = _runs_to_markdown(_iter_paragraph_items(para)) or text

            if level > 0:
                # Strip outer **...** wrapping that headings may have inherited
                clean = inline.strip()
                if clean.startswith("**") and clean.endswith("**"):
                    clean = clean[2:-2]
                if prev_was_list:
                    lines.append("")
                prev_was_list = False
                lines.append(f"{'#' * level} {clean}")
                lines.append("")
            else:
                list_info = _get_list_info(para, numbering_map)
                if list_info:
                    list_type, ilvl, num_id = list_info
                    indent = "    " * ilvl
                    if list_type == "ordered":
                        counter_key = f"{num_id}-{ilvl}"
                        ol_counters[counter_key] = ol_counters.get(counter_key, 0) + 1
                        prefix = f"{indent}{ol_counters[counter_key]}. "
                    else:
                        prefix = f"{indent}- "
                    if not prev_was_list and lines and lines[-1] != "":
                        lines.append("")
                    prev_was_list = True
                    lines.append(prefix + inline)
                else:
                    if prev_was_list:
                        lines.append("")
                    prev_was_list = False
                    # Reset ordered list counters when a list is interrupted
                    ol_counters.clear()
                    lines.append(inline)
                    lines.append("")

        elif tag == "tbl":
            flush_code_buf()
            if prev_was_list:
                lines.append("")
            prev_was_list = False
            ol_counters.clear()
            table = Table(child, doc)
            if lines and lines[-1] != "":
                lines.append("")
            lines.append(
                _table_to_html(table, doc, image_counter, images, content_width)
            )
            lines.append("")

    # Strip trailing blank lines
    flush_code_buf()
    while lines and lines[-1] == "":
        lines.pop()

    return "\n".join(lines), images


@default_registry.register
class DocxConverter(BaseConverter):
    supported_extensions = ["docx"]
    supported_mimetypes = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    def convert_file(self, file_path: Path, **kwargs) -> ConvertResult:
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError(
                "DOCX conversion requires python-docx: pip install paddleocr[doc2md]"
            )

        doc = Document(file_path)
        md_text, images = _convert_body(doc)

        return ConvertResult(
            markdown=md_text,
            images=images,
            metadata={
                "format": "DOCX",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            },
        )
