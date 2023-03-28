# Copyright (c) 2022 PaddlePaddle Authors. All Rights Reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
"""
This code is refer from: https://github.com/weizwx/html2docx/blob/master/htmldocx/h2d.py
"""

import re
import docx
from docx import Document
from bs4 import BeautifulSoup
from html.parser import HTMLParser


def get_table_rows(table_soup):
    table_row_selectors = [
        'table > tr', 'table > thead > tr', 'table > tbody > tr',
        'table > tfoot > tr'
    ]
    # If there's a header, body, footer or direct child tr tags, add row dimensions from there
    return table_soup.select(', '.join(table_row_selectors), recursive=False)


def get_table_columns(row):
    # Get all columns for the specified row tag.
    return row.find_all(['th', 'td'], recursive=False) if row else []


def get_table_dimensions(table_soup):
    # Get rows for the table
    rows = get_table_rows(table_soup)
    # Table is either empty or has non-direct children between table and tr tags
    # Thus the row dimensions and column dimensions are assumed to be 0

    cols = get_table_columns(rows[0]) if rows else []
    # Add colspan calculation column number
    col_count = 0
    for col in cols:
        colspan = col.attrs.get('colspan', 1)
        col_count += int(colspan)

    return rows, col_count


def get_cell_html(soup):
    # Returns string of td element with opening and closing <td> tags removed
    # Cannot use find_all as it only finds element tags and does not find text which
    # is not inside an element
    return ' '.join([str(i) for i in soup.contents])


def delete_paragraph(paragraph):
    # https://github.com/python-openxml/python-docx/issues/33#issuecomment-77661907
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def remove_whitespace(string, leading=False, trailing=False):
    """Remove white space from a string.
    Args:
        string(str): The string to remove white space from.
        leading(bool, optional): Remove leading new lines when True.
        trailing(bool, optional): Remove trailing new lines when False.
    Returns:
        str: The input string with new line characters removed and white space squashed.
    Examples:
        Single or multiple new line characters are replaced with space.
            >>> remove_whitespace("abc\\ndef")
            'abc def'
            >>> remove_whitespace("abc\\n\\n\\ndef")
            'abc def'
        New line characters surrounded by white space are replaced with a single space.
            >>> remove_whitespace("abc \\n \\n \\n def")
            'abc def'
            >>> remove_whitespace("abc  \\n  \\n  \\n  def")
            'abc def'
        Leading and trailing new lines are replaced with a single space.
            >>> remove_whitespace("\\nabc")
            ' abc'
            >>> remove_whitespace("  \\n  abc")
            ' abc'
            >>> remove_whitespace("abc\\n")
            'abc '
            >>> remove_whitespace("abc  \\n  ")
            'abc '
        Use ``leading=True`` to remove leading new line characters, including any surrounding
        white space:
            >>> remove_whitespace("\\nabc", leading=True)
            'abc'
            >>> remove_whitespace("  \\n  abc", leading=True)
            'abc'
        Use ``trailing=True`` to remove trailing new line characters, including any surrounding
        white space:
            >>> remove_whitespace("abc  \\n  ", trailing=True)
            'abc'
    """
    # Remove any leading new line characters along with any surrounding white space
    if leading:
        string = re.sub(r'^\s*\n+\s*', '', string)

    # Remove any trailing new line characters along with any surrounding white space
    if trailing:
        string = re.sub(r'\s*\n+\s*$', '', string)

    # Replace new line characters and absorb any surrounding space.
    string = re.sub(r'\s*\n\s*', ' ', string)
    # TODO need some way to get rid of extra spaces in e.g. text <span>   </span>  text
    return re.sub(r'\s+', ' ', string)


font_styles = {
    'b': 'bold',
    'strong': 'bold',
    'em': 'italic',
    'i': 'italic',
    'u': 'underline',
    's': 'strike',
    'sup': 'superscript',
    'sub': 'subscript',
    'th': 'bold',
}

font_names = {
    'code': 'Courier',
    'pre': 'Courier',
}


class HtmlToDocx(HTMLParser):
    def __init__(self):
        super().__init__()
        self.options = {
            'fix-html': True,
            'images': True,
            'tables': True,
            'styles': True,
        }
        self.table_row_selectors = [
            'table > tr', 'table > thead > tr', 'table > tbody > tr',
            'table > tfoot > tr'
        ]
        self.table_style = None
        self.paragraph_style = None

    def set_initial_attrs(self, document=None):
        self.tags = {
            'span': [],
            'list': [],
        }
        if document:
            self.doc = document
        else:
            self.doc = Document()
        self.bs = self.options[
            'fix-html']  # whether or not to clean with BeautifulSoup
        self.document = self.doc
        self.include_tables = True  #TODO add this option back in?
        self.include_images = self.options['images']
        self.include_styles = self.options['styles']
        self.paragraph = None
        self.skip = False
        self.skip_tag = None
        self.instances_to_skip = 0

    def copy_settings_from(self, other):
        """Copy settings from another instance of HtmlToDocx"""
        self.table_style = other.table_style
        self.paragraph_style = other.paragraph_style

    def ignore_nested_tables(self, tables_soup):
        """
        Returns array containing only the highest level tables
        Operates on the assumption that bs4 returns child elements immediately after
        the parent element in `find_all`. If this changes in the future, this method will need to be updated
        :return:
        """
        new_tables = []
        nest = 0
        for table in tables_soup:
            if nest:
                nest -= 1
                continue
            new_tables.append(table)
            nest = len(table.find_all('table'))
        return new_tables

    def get_tables(self):
        if not hasattr(self, 'soup'):
            self.include_tables = False
            return
            # find other way to do it, or require this dependency?
        self.tables = self.ignore_nested_tables(self.soup.find_all('table'))
        self.table_no = 0

    def run_process(self, html):
        if self.bs and BeautifulSoup:
            self.soup = BeautifulSoup(html, 'html.parser')
            html = str(self.soup)
        if self.include_tables:
            self.get_tables()
        self.feed(html)

    def add_html_to_cell(self, html, cell):
        if not isinstance(cell, docx.table._Cell):
            raise ValueError('Second argument needs to be a %s' %
                             docx.table._Cell)
        unwanted_paragraph = cell.paragraphs[0]
        if unwanted_paragraph.text == "":
            delete_paragraph(unwanted_paragraph)
        self.set_initial_attrs(cell)
        self.run_process(html)
        # cells must end with a paragraph or will get message about corrupt file
        # https://stackoverflow.com/a/29287121
        if not self.doc.paragraphs:
            self.doc.add_paragraph('')

    def apply_paragraph_style(self, style=None):
        try:
            if style:
                self.paragraph.style = style
            elif self.paragraph_style:
                self.paragraph.style = self.paragraph_style
        except KeyError as e:
            raise ValueError(
                f"Unable to apply style {self.paragraph_style}.") from e

    def handle_table(self, html, doc):
        """
        To handle nested tables, we will parse tables manually as follows:
        Get table soup
        Create docx table
        Iterate over soup and fill docx table with new instances of this parser
        Tell HTMLParser to ignore any tags until the corresponding closing table tag
        """
        table_soup = BeautifulSoup(html, 'html.parser')
        rows, cols_len = get_table_dimensions(table_soup)
        table = doc.add_table(len(rows), cols_len)
        table.style = doc.styles['Table Grid']

        cell_row = 0
        for index, row in enumerate(rows):
            cols = get_table_columns(row)
            cell_col = 0
            for col in cols:
                colspan = int(col.attrs.get('colspan', 1))
                rowspan = int(col.attrs.get('rowspan', 1))

                cell_html = get_cell_html(col)
                if col.name == 'th':
                    cell_html = "<b>%s</b>" % cell_html

                docx_cell = table.cell(cell_row, cell_col)

                while docx_cell.text != '':  # Skip the merged cell
                    cell_col += 1
                    docx_cell = table.cell(cell_row, cell_col)

                cell_to_merge = table.cell(cell_row + rowspan - 1,
                                           cell_col + colspan - 1)
                if docx_cell != cell_to_merge:
                    docx_cell.merge(cell_to_merge)

                child_parser = HtmlToDocx()
                child_parser.copy_settings_from(self)
                child_parser.add_html_to_cell(cell_html or ' ', docx_cell)

                cell_col += colspan
            cell_row += 1

    def handle_data(self, data):
        if self.skip:
            return

        # Only remove white space if we're not in a pre block.
        if 'pre' not in self.tags:
            # remove leading and trailing whitespace in all instances
            data = remove_whitespace(data, True, True)

        if not self.paragraph:
            self.paragraph = self.doc.add_paragraph()
            self.apply_paragraph_style()

        # There can only be one nested link in a valid html document
        # You cannot have interactive content in an A tag, this includes links
        # https://html.spec.whatwg.org/#interactive-content
        link = self.tags.get('a')
        if link:
            self.handle_link(link['href'], data)
        else:
            # If there's a link, dont put the data directly in the run
            self.run = self.paragraph.add_run(data)
            spans = self.tags['span']
            for span in spans:
                if 'style' in span:
                    style = self.parse_dict_string(span['style'])
                    self.add_styles_to_run(style)

            # add font style and name
            for tag in self.tags:
                if tag in font_styles:
                    font_style = font_styles[tag]
                    setattr(self.run.font, font_style, True)

                if tag in font_names:
                    font_name = font_names[tag]
                    self.run.font.name = font_name
