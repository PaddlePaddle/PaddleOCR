import os
import sys
import cv2
from docx import Document
from paddleocr import PPStructure

current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(os.path.abspath(os.path.join(current_dir, "..")))

from ppstructure.recovery.recovery_to_doc import sorted_layout_boxes, convert_info_docx


def test_double_column_structure(tmp_path):
    """
    Test document structure analysis and docx generation for double-column layout.
    Validates layout detection, column separation and document conversion.
    """
    img_path = os.path.join(current_dir, "./test_files/double_column.png")
    assert os.path.exists(img_path), "Required test image file not found"

    # Initialize test image
    img = cv2.imread(img_path)
    assert img is not None, "Image loading failed"

    # Initialize PPStructure engine for layout analysis
    engine = PPStructure(show_log=True)

    # Perform structural analysis
    result = engine(img)
    assert (
        result is not None and len(result) > 0
    ), "Structure analysis produced no results"

    # Process layout boxes in reading order
    img_h, img_w = img.shape[:2]
    sorted_results = sorted_layout_boxes(result, img_w)

    # Validate double-column layout detection
    double_column_boxes = []
    for item in sorted_results:
        if item.get("layout") == "double":
            double_column_boxes.append(item)

    assert len(double_column_boxes) >= 2, "Double-column layout not properly detected"

    # Validate column distribution
    left_column = []
    right_column = []
    for box in double_column_boxes:
        box_center = (box["bbox"][0] + box["bbox"][2]) / 2
        if box_center < img_w / 2:
            left_column.append(box)
        else:
            right_column.append(box)

    assert len(left_column) > 0, "Left column content not detected"
    assert len(right_column) > 0, "Right column content not detected"

    # Configure output directory
    output_dir = str(tmp_path / "double_column_test")
    os.makedirs(output_dir, exist_ok=True)

    # Process document conversion
    img_name = "test_double_column"
    convert_info_docx(img, sorted_results, output_dir, img_name)

    # Validate output document existence
    docx_path = os.path.join(output_dir, f"{img_name}_ocr.docx")
    assert os.path.exists(docx_path), "Document generation failed"

    # Verify document content
    doc = Document(docx_path)
    assert len(doc.paragraphs) > 0, "Generated document contains no content"


def test_single_column_structure(tmp_path):
    """
    Test document structure analysis and docx generation for single-column layout.
    Validates layout detection, width ratio analysis and document conversion.
    """
    img_path = os.path.join(current_dir, "./test_files/single_column.jpg")
    assert os.path.exists(img_path), f"Test image {img_path} not found"

    img = cv2.imread(img_path)
    assert img is not None, f"Failed to load image {img_path}"

    engine = PPStructure(show_log=True)
    result = engine(img)
    assert result is not None and len(result) > 0, "Layout analysis result is empty"

    img_h, img_w = img.shape[:2]
    sorted_results = sorted_layout_boxes(result, img_w)

    # Check layout assignment is correct
    single_column_boxes = []
    for item in sorted_results:
        box_width = item["bbox"][2] - item["bbox"][0]
        width_ratio = box_width / img_w
        # For text boxes that span >60% of page width, verify they are marked as single column
        if width_ratio > 0.6 and item["type"] == "text":
            assert (
                item.get("layout") == "single"
            ), f"Wide text box ({width_ratio:.2f} of page width) not marked as single column"
            single_column_boxes.append(item)

    assert len(single_column_boxes) > 0, "No single column text boxes detected"

    # Use temporary directory for output
    output_dir = str(tmp_path / "single_column_test")
    os.makedirs(output_dir, exist_ok=True)

    img_name = "test_single_column"
    convert_info_docx(img, sorted_results, output_dir, img_name)

    # Verify output document
    docx_path = os.path.join(output_dir, f"{img_name}_ocr.docx")
    assert os.path.exists(docx_path), "Document not generated"

    # Validate document content
    doc = Document(docx_path)
    assert len(doc.paragraphs) > 0, "Generated document is empty"
